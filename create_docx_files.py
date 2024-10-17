import os
import pandas as pd
from docx import Document

# Функция для чтения существующего .docx файла и получения списка артикулов
def get_existing_articles_from_docx(docx_file):
    if os.path.exists(docx_file):
        doc = Document(docx_file)
        existing_articles = []
        for row in doc.tables[0].rows[1:]:  # Пропускаем заголовок таблицы
            article = row.cells[2].text.strip()  # Артикул находится в 3-й колонке
            existing_articles.append(article)
        return existing_articles
    return []

# Функция для создания .docx файла на основе Excel файла с проверкой на дубликаты
def save_unique_products_to_docx(file_name, output_docx_file):
    excel_file = f"downloads/{file_name}"  # Путь к Excel файлу
    try:
        if os.path.exists(excel_file):
            print(f"Чтение файла: {excel_file}")
            # Читаем данные из Excel файла
            df = pd.read_excel(excel_file)

            print(f"Количество строк в файле: {len(df)}")
            if df.empty:
                print("Файл пуст.")
                return None

            # Проверяем наличие необходимых колонок
            if 'Название' not in df.columns or 'Артикул' not in df.columns:
                print("Файл не содержит необходимых колонок 'Название' или 'Артикул'.")
                return None

            # Получаем артикулы из существующего .docx файла
            existing_articles = get_existing_articles_from_docx(output_docx_file)

            # Убираем дубликаты по колонке "Артикул", и проверяем, нет ли их в существующем .docx файле
            new_products = df[['Название', 'Артикул']].drop_duplicates(subset=['Артикул'])
            unique_products = new_products[~new_products['Артикул'].isin(existing_articles)]
            print(f"Количество новых товаров для добавления: {len(unique_products)}")

            # Если нет новых товаров, выходим
            if unique_products.empty:
                print("Нет новых товаров для добавления.")
                return None

            # Создаем или открываем существующий .docx файл
            if os.path.exists(output_docx_file):
                doc = Document(output_docx_file)
                table = doc.tables[0]  # Предполагаем, что таблица одна и уже существует
            else:
                doc = Document()
                doc.add_heading(f'Список уникальных товаров из {file_name}', 0)
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = '№'
                hdr_cells[1].text = 'Название'
                hdr_cells[2].text = 'Артикул'

            # Заполняем таблицу новыми товарами
            existing_rows = len(table.rows)  # Уже существующие строки
            for idx, row in enumerate(unique_products.itertuples(), 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(existing_rows + idx - 1)
                row_cells[1].text = str(row.Название)
                row_cells[2].text = str(row.Артикул)

            # Сохраняем документ
            doc.save(output_docx_file)
            print(f"Документ обновлен: {output_docx_file}")
            return output_docx_file
        else:
            print(f"Файл {excel_file} не найден.")
            return None
    except Exception as e:
        print(f"Ошибка при обработке файла {file_name}: {e}")
        return None

# Список файлов и соответствующих имен для сохранения
files = [
    ("parsed_data.xlsx", "downloads/parsed_data_products.docx"),
    ("parsed_data2.xlsx", "downloads/parsed_data2_products.docx"),
    ("parsed_data_autoopt.xlsx", "downloads/parsed_data_autoopt_products.docx"),
]

# Обработка каждого файла
for excel_file, docx_file in files:
    save_unique_products_to_docx(excel_file, docx_file)
